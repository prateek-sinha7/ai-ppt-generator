"""
DOCX Export Service — generates Word documents from Slide_JSON using python-docx.
"""
import io
from typing import Any, Dict, List, Optional

import structlog
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = structlog.get_logger(__name__)


# ---------------------------------------------------------------------------
# Theme color palettes (RGB tuples for python-docx RGBColor)
# ---------------------------------------------------------------------------

THEME_PALETTES: Dict[str, Dict[str, Any]] = {
    "ocean-depths": {
        "primary": (26, 35, 50),
        "secondary": (45, 139, 139),
        "accent": (168, 218, 220),
        "text": (26, 35, 50),
        "text_light": (107, 114, 128),
        "background": (241, 250, 238),
        "header_bar": (26, 35, 50),
    },
    "sunset-boulevard": {
        "primary": (231, 111, 81),
        "secondary": (244, 162, 97),
        "accent": (233, 196, 106),
        "text": (38, 70, 83),
        "text_light": (107, 114, 128),
        "background": (255, 255, 255),
        "header_bar": (38, 70, 83),
    },
    "forest-canopy": {
        "primary": (45, 74, 43),
        "secondary": (125, 132, 113),
        "accent": (164, 172, 134),
        "text": (45, 74, 43),
        "text_light": (107, 114, 128),
        "background": (250, 249, 246),
        "header_bar": (45, 74, 43),
    },
    "modern-minimalist": {
        "primary": (54, 69, 79),
        "secondary": (112, 128, 144),
        "accent": (211, 211, 211),
        "text": (54, 69, 79),
        "text_light": (112, 128, 144),
        "background": (255, 255, 255),
        "header_bar": (54, 69, 79),
    },
    "golden-hour": {
        "primary": (244, 169, 0),
        "secondary": (193, 102, 107),
        "accent": (212, 184, 150),
        "text": (74, 64, 58),
        "text_light": (107, 114, 128),
        "background": (255, 255, 255),
        "header_bar": (74, 64, 58),
    },
    "arctic-frost": {
        "primary": (74, 111, 165),
        "secondary": (192, 192, 192),
        "accent": (212, 228, 247),
        "text": (44, 62, 80),
        "text_light": (107, 114, 128),
        "background": (250, 250, 250),
        "header_bar": (74, 111, 165),
    },
    "desert-rose": {
        "primary": (212, 165, 165),
        "secondary": (184, 125, 109),
        "accent": (232, 213, 196),
        "text": (93, 46, 70),
        "text_light": (107, 114, 128),
        "background": (255, 255, 255),
        "header_bar": (93, 46, 70),
    },
    "tech-innovation": {
        "primary": (0, 102, 255),
        "secondary": (0, 255, 255),
        "accent": (0, 204, 204),
        "text": (255, 255, 255),
        "text_light": (156, 163, 175),
        "background": (30, 30, 30),
        "header_bar": (0, 102, 255),
    },
    "botanical-garden": {
        "primary": (74, 124, 89),
        "secondary": (249, 166, 32),
        "accent": (183, 71, 42),
        "text": (58, 58, 58),
        "text_light": (107, 114, 128),
        "background": (245, 243, 237),
        "header_bar": (74, 124, 89),
    },
    "midnight-galaxy": {
        "primary": (74, 78, 143),
        "secondary": (164, 144, 194),
        "accent": (230, 230, 250),
        "text": (230, 230, 250),
        "text_light": (156, 163, 175),
        "background": (43, 30, 62),
        "header_bar": (74, 78, 143),
    },
}

_DEFAULT_THEME = "ocean-depths"


def _resolve_theme(theme: str) -> Dict[str, Any]:
    """Resolve a theme name to its color palette, falling back to ocean-depths."""
    key = theme.lower().replace("_", "-")
    return THEME_PALETTES.get(key, THEME_PALETTES[_DEFAULT_THEME])


def _rgb(color_tuple: tuple) -> RGBColor:
    """Convert an (r, g, b) tuple to a python-docx RGBColor."""
    return RGBColor(*color_tuple)


def _extract_slides(slides_data: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Extract the slides list from slides_data (dict with 'slides' key or list)."""
    if isinstance(slides_data, list):
        return slides_data
    if isinstance(slides_data, dict):
        slides = slides_data.get("slides", slides_data.get("slide_json", []))
        if isinstance(slides, list):
            return slides
    return []


# ---------------------------------------------------------------------------
# Table rendering helper
# ---------------------------------------------------------------------------

def _add_table(doc: Document, table_data: Dict[str, Any], palette: Dict[str, Any]) -> None:
    """Add a formatted table to the document."""
    headers = table_data.get("headers", [])
    rows = table_data.get("rows", [])
    if not headers and not rows:
        return

    col_count = len(headers) if headers else (len(rows[0]) if rows else 0)
    if col_count == 0:
        return

    row_count = (1 if headers else 0) + len(rows)
    tbl = doc.add_table(rows=row_count, cols=col_count)
    tbl.style = "Table Grid"

    # Header row
    start_row = 0
    if headers:
        header_bg = palette.get("header_bar", palette["primary"])
        for ci, header_text in enumerate(headers):
            cell = tbl.rows[0].cells[ci]
            cell.text = str(header_text)
            # Style header cell
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(str(header_text))
            if not cell.paragraphs[0].runs:
                cell.paragraphs[0].clear()
                run = cell.paragraphs[0].add_run(str(header_text))
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
            # Set cell background
            _set_cell_bg(cell, header_bg)
        start_row = 1

    # Data rows
    for ri, row in enumerate(rows):
        row_cells = tbl.rows[start_row + ri].cells
        if isinstance(row, dict):
            for ci, h in enumerate(headers):
                if ci < len(row_cells):
                    cell = row_cells[ci]
                    cell.text = str(row.get(h, ""))
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = _rgb(palette["text"])
        elif isinstance(row, (list, tuple)):
            for ci, val in enumerate(row):
                if ci < len(row_cells):
                    cell = row_cells[ci]
                    cell.text = str(val)
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = _rgb(palette["text"])


def _set_cell_bg(cell: Any, color_tuple: tuple) -> None:
    """Set the background (shading) color of a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "{:02x}{:02x}{:02x}".format(*color_tuple))
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def generate_docx(
    slides_data: Dict[str, Any],
    theme: str = "ocean-depths",
    design_spec: Optional[Dict] = None,
) -> bytes:
    """Generate a DOCX document from Slide_JSON.

    Args:
        slides_data: Dict with a ``"slides"`` key (list of slide dicts), or a
            plain list of slide dicts.
        theme: Theme identifier (e.g. ``"ocean-depths"``).
        design_spec: Optional design specification dict (currently unused but
            reserved for future refinement).

    Returns:
        DOCX file content as ``bytes``.
    """
    palette = _resolve_theme(theme)
    slides = _extract_slides(slides_data)

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = _rgb(palette["text"])

    for idx, slide in enumerate(slides):
        try:
            _render_slide(doc, slide, palette)
        except Exception:
            logger.warning("docx_export.slide_render_failed", slide_index=idx, exc_info=True)

        # Page break after every slide except the last
        if idx < len(slides) - 1:
            doc.add_page_break()

    if not slides:
        doc.add_paragraph("No slides to export.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_slide(doc: Document, slide: Dict[str, Any], palette: Dict[str, Any]) -> None:
    """Render a single slide into the Word document."""
    title = slide.get("title", "")
    subtitle = slide.get("subtitle", "")
    bullets = slide.get("bullets", slide.get("content", []))
    table_data = slide.get("table_data", slide.get("tableData", None))

    # Title as Heading 1 with theme primary color
    if title:
        heading = doc.add_heading(level=1)
        run = heading.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = _rgb(palette["primary"])

    # Subtitle
    if subtitle:
        sub_para = doc.add_paragraph()
        run = sub_para.add_run(subtitle)
        run.font.size = Pt(14)
        run.font.color.rgb = _rgb(palette.get("secondary", palette["primary"]))
        run.italic = True

    # Bullets
    if isinstance(bullets, list) and bullets:
        for bullet in bullets:
            text = bullet if isinstance(bullet, str) else str(bullet)
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(text)
            run.font.size = Pt(11)
            run.font.color.rgb = _rgb(palette["text"])

    # Table
    if table_data and isinstance(table_data, dict):
        doc.add_paragraph()  # spacer
        _add_table(doc, table_data, palette)
